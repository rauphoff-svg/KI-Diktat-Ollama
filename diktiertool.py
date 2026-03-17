#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
KI-Diktat-Ollama – Vollständig lokale Diktierlösung
====================================================
Push-to-Talk Aufnahme -> FasterWhisper (lokal) -> Ollama-Korrektur -> DOCX/TXT

Kein Datenaustausch mit externen Servern. Geeignet für Anwaltsdiktat
unter strengen Datenschutzanforderungen (DSGVO, Anwaltsgeheimnis).

Backends:
  Transkription: FasterWhisper (lokal, Modell konfigurierbar)
  Korrektur:     Ollama (lokal) – mistral-small:latest oder gemma3:latest

Aenderungsprotokoll v1.1:
  - Aufnahmesteuerung: pynput Key-Objekte statt String-Vergleich (Bugfix)
  - _KEY_MAP + resolve_key() wie in KI-Diktat-Voxtral-Claude
  - recording_worker Thread wie im Original
  - on_key_press / on_key_release als Klassenmethoden

Aenderungsprotokoll v1:
  - Aufbau auf Basis von KI-Diktat-Voxtral-Claude v4.2
  - Transkription: Voxtral -> FasterWhisper (lokal)
  - Korrektur: Mistral/Claude -> Ollama (lokal)
  - Modellwechsel: mistral-small <-> gemma3

Abhaengigkeiten:
  pip install faster-whisper numpy sounddevice soundfile pyperclip pynput pyyaml python-docx
  Ollama muss lokal laufen (https://ollama.com)
"""

import os
import sys
import json
import time
import queue
import threading
import re
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Optional, TypedDict

import yaml
import numpy as np
import sounddevice as sd
import pyperclip
from pynput import keyboard

# Windows-Terminal: stdout auf UTF-8 umstellen
if sys.stdout and hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

# FasterWhisper
try:
    from faster_whisper import WhisperModel
    _HAS_FASTER_WHISPER = True
except ImportError:
    _HAS_FASTER_WHISPER = False

# python-docx (optional)
try:
    from docx import Document
    from docx.shared import Pt
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


# =========================================================================
# TypedDicts
# =========================================================================

class AudioConfig(TypedDict):
    sample_rate: int
    channels: int
    blocksize: int
    silence_threshold_rms: float


class TranscriptionConfig(TypedDict):
    model: str
    language: str
    device: str
    compute_type: str


class CorrectionConfig(TypedDict):
    model: str
    ollama_base_url: str
    max_tokens: int
    temperature: float
    paragraph_wise: bool


class OutputConfig(TypedDict):
    format: str
    directory: str
    docx_font: str
    docx_font_size_pt: int
    docx_template: Optional[str]


class ControlsConfig(TypedDict):
    push_to_talk_key: str
    pause_key: str
    reject_segment_key: str
    segment_review: bool


class AppConfig(TypedDict):
    audio: AudioConfig
    transcription: TranscriptionConfig
    correction: CorrectionConfig
    output: OutputConfig
    controls: ControlsConfig


# =========================================================================
# Konfiguration
# =========================================================================

_DEFAULTS: dict = {
    "audio": {
        "sample_rate": 16000,
        "channels": 1,
        "blocksize": 2048,
        "silence_threshold_rms": 0.005,
    },
    "transcription": {
        "model": "medium",
        "language": "de",
        "device": "cpu",
        "compute_type": "int8",
    },
    "correction": {
        "model": "mistral-small:latest",
        "ollama_base_url": "http://localhost:11434",
        "max_tokens": 4096,
        "temperature": 0.1,
        "paragraph_wise": True,
    },
    "output": {
        "format": "txt",
        "directory": ".",
        "docx_font": "Arial",
        "docx_font_size_pt": 11,
        "docx_template": None,
    },
    "controls": {
        "push_to_talk_key": "ctrl_r",
        "pause_key": "f10",
        "reject_segment_key": "f9",
        "segment_review": True,
    },
}

AVAILABLE_MODELS = ["mistral-small:latest", "gemma3:latest"]


def _deep_merge(base: dict, override: dict) -> dict:
    result = dict(base)
    for k, v in override.items():
        if k in result and isinstance(result[k], dict) and isinstance(v, dict):
            result[k] = _deep_merge(result[k], v)
        else:
            result[k] = v
    return result


def load_config() -> AppConfig:
    candidates = [
        Path(__file__).parent / "config.yaml",
        Path.cwd() / "config.yaml",
    ]
    cfg_path = next((p for p in candidates if p.is_file()), candidates[0])
    if cfg_path.is_file():
        with open(cfg_path, encoding="utf-8") as f:
            user_cfg = yaml.safe_load(f) or {}
        return _deep_merge(_DEFAULTS, user_cfg)  # type: ignore
    return dict(_DEFAULTS)  # type: ignore


# =========================================================================
# Tastenbelegung – pynput Key-Objekte
# =========================================================================

_KEY_MAP = {
    "ctrl_r":  keyboard.Key.ctrl_r,
    "ctrl_l":  keyboard.Key.ctrl_l,
    "alt_r":   keyboard.Key.alt_r,
    "alt_l":   keyboard.Key.alt_l,
    "shift_r": keyboard.Key.shift_r,
    "f1":  keyboard.Key.f1,  "f2":  keyboard.Key.f2,
    "f3":  keyboard.Key.f3,  "f4":  keyboard.Key.f4,
    "f5":  keyboard.Key.f5,  "f6":  keyboard.Key.f6,
    "f7":  keyboard.Key.f7,  "f8":  keyboard.Key.f8,
    "f9":  keyboard.Key.f9,  "f10": keyboard.Key.f10,
    "f11": keyboard.Key.f11, "f12": keyboard.Key.f12,
}


def resolve_key(name: str):
    """Löst einen Tastennamen in ein pynput-Key-Objekt auf."""
    return _KEY_MAP.get(name.lower(), keyboard.Key.ctrl_r)


# =========================================================================
# Audio-Hilfsfunktionen
# =========================================================================

def rms(audio_data: np.ndarray) -> float:
    if audio_data.size == 0:
        return 0.0
    val = float(np.sqrt(np.mean(np.square(audio_data.astype(np.float32)))))
    return max(0.0, val - 0.003)


def level_bar(rms_val: float, width: int = 30) -> str:
    filled = int(min(rms_val / 0.15, 1.0) * width)
    return "[" + "█" * filled + "░" * (width - filled) + f"] {rms_val:.3f}"


# =========================================================================
# Korrektur-Prompt
# =========================================================================

CORRECTION_SYSTEM_PROMPT = """\
Du bist ein juristischer Lektor. Deine Aufgabe ist ausschliesslich die \
sprachliche und formale Korrektur des diktierten Textes.

WICHTIGSTE REGEL: Du darfst NICHTS inhaltlich hinzufuegen!
- KEINE Paragraphen, Normen oder Gesetzeszitate ergaenzen, die nicht diktiert wurden
- KEINE juristischen Argumente, Begruendungen oder Erlaeuterungen einfuegen
- KEINE Sachverhaltselemente, Daten oder Fakten hinzudichten
- Der diktierte Inhalt ist die einzige Quelle – du korrigierst NUR Sprache und Form

Korrekturregeln:
1. Ausgesprochene Satzzeichen umwandeln:
   - "Komma" -> ,
   - "Punkt" / "Satzpunkt" -> .
   - "Doppelpunkt" -> :
   - "Semikolon" -> ;
   - "Gedankenstrich" -> \u2013 (mit Leerzeichen davor und danach)
   - "Bindestrich" -> -
   - "Schraegstrich" -> /
   - "Ausrufezeichen" -> !
   - "Fragezeichen" -> ?
   - "Klammer auf" -> (
   - "Klammer zu" -> )
   - "Anfuehrungszeichen auf" -> \u201e
   - "Anfuehrungszeichen zu" / "Anfuehrungszeichen Ende" -> \u201c
2. Normzitate: "Paragraph 823 Absatz 2 Satz 1 BGB" -> "\u00a7 823 Abs. 2 S. 1 BGB"
3. Abkuerzungen: Nur standardisierte juristische Abkuerzungen (vgl., i.V.m., i.S.d.)
4. Stil: Direkt und verstaendlich, Aktivsaetze bevorzugen, kein Kanzleideutsch
5. Interpunktion: Korrekte deutsche Zeichensetzung
6. Grammatik: Korrekte Deklination, Konjugation, Kongruenz
7. Rechtschreibung: Korrekte deutsche Rechtschreibung (Duden)
8. Absaetze: Erhalte Absatzstruktur exakt wie diktiert
9. Steuerkommandos:
   - "Seitenumbruch" oder "neue Seite" -> [SEITENUMBRUCH]
   - "Aufzaehlung:" -> echte Aufzaehlung mit Spiegelstrichen
   - "neuer Absatz" -> Absatzumbruch
10. Waehrung: Immer "EUR"
11. Prozent: Immer Zahl + Leerzeichen + % (z.B. "50 %")
12. Doppelwoerter entfernen (z.B. "der der Klaeger" -> "der Klaeger")

Antworte NUR mit dem korrigierten Text, ohne Erklaerungen oder Kommentare.\
"""

CORRECTION_PARAGRAPH_PREFIX = """\
Vorheriger Absatz (nur als Kontext, nicht nochmals ausgeben):
{context}

Zu korrigierender Absatz:
{text}"""

CORRECTION_FULL_PREFIX = "{text}"


# =========================================================================
# Ollama-Hilfsfunktionen
# =========================================================================

def _call_ollama(
    base_url: str,
    model: str,
    system: str,
    user: str,
    temperature: float = 0.1,
    max_tokens: int = 4096,
    timeout: int = 120,
) -> str:
    payload = json.dumps({
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "stream": False,
        "options": {
            "temperature": temperature,
            "num_predict": max_tokens,
        },
    }).encode("utf-8")
    req = urllib.request.Request(
        f"{base_url}/api/chat",
        data=payload,
        headers={"Content-Type": "application/json"},
    )
    with urllib.request.urlopen(req, timeout=timeout) as r:
        resp = json.loads(r.read().decode("utf-8"))
    return resp["message"]["content"].strip()


def _check_ollama(base_url: str, timeout: int = 3) -> bool:
    try:
        urllib.request.urlopen(f"{base_url}/api/tags", timeout=timeout)
        return True
    except Exception:
        return False


def _list_ollama_models(base_url: str) -> list[str]:
    try:
        with urllib.request.urlopen(f"{base_url}/api/tags", timeout=3) as r:
            data = json.loads(r.read().decode("utf-8"))
        return [m["name"] for m in data.get("models", [])]
    except Exception:
        return []


# =========================================================================
# Geräteauswahl
# =========================================================================

def choose_input_device() -> int:
    devices = sd.query_devices()
    inputs = [(i, d) for i, d in enumerate(devices) if d["max_input_channels"] > 0]
    if not inputs:
        raise RuntimeError("Kein Mikrofon gefunden.")
    print("\n  Verfügbare Mikrofone:")
    for i, d in inputs:
        print(f"    [{i}] {d['name']}")
    default_idx = sd.default.device[0]
    default_name = devices[default_idx]["name"] if default_idx is not None else ""
    print(f"\n  Standard: [{default_idx}] {default_name}")
    raw = input("  Auswahl (Enter = Standard): ").strip()
    if not raw:
        return default_idx
    idx = int(raw)
    if idx not in [i for i, _ in inputs]:
        raise ValueError(f"Ungültige Auswahl: {idx}")
    return idx


# =========================================================================
# Transkription mit FasterWhisper
# =========================================================================

class Transcriber:
    """Wrapper um FasterWhisper mit Lazy-Loading."""

    def __init__(self, cfg: AppConfig):
        self.cfg = cfg
        self._model = None

    def _load(self):
        if self._model is not None:
            return
        t_cfg = self.cfg["transcription"]
        print(f"  Lade Whisper-Modell '{t_cfg['model']}' "
              f"({t_cfg['device']}/{t_cfg['compute_type']}) …", end=" ", flush=True)
        self._model = WhisperModel(
            t_cfg["model"],
            device=t_cfg["device"],
            compute_type=t_cfg["compute_type"],
        )
        print("OK")

    def transcribe(self, audio: np.ndarray) -> str:
        """Transkribiert ein float32-Audio-Array."""
        self._load()
        lang = self.cfg["transcription"]["language"]
        segments, _ = self._model.transcribe(
            audio,
            language=lang,
            beam_size=5,
            vad_filter=True,
        )
        return " ".join(seg.text.strip() for seg in segments).strip()


# =========================================================================
# DictationSession
# =========================================================================

class DictationSession:
    """Kapselt eine vollständige Diktat-Sitzung."""

    def __init__(self, cfg: AppConfig, transcriber: Transcriber):
        self.cfg = cfg
        self.transcriber = transcriber

        # Audio
        self.audio_queue: queue.Queue = queue.Queue()
        self.sample_rate: int = cfg["audio"]["sample_rate"]
        self.channels: int = cfg["audio"]["channels"]
        self.blocksize: int = cfg["audio"]["blocksize"]

        # Zustand
        self.recording = False
        self.paused = False
        self.segment_chunks: list[np.ndarray] = []
        self.text_blocks: list[str] = []
        self.finished = False
        self.aborted = False
        self.lock = threading.Lock()

        # Verarbeitungs-Lock: verhindert parallele Segment-Verarbeitung
        self._processing_lock = threading.Lock()

        # Tasten – als pynput Key-Objekte auflösen
        ctrl = cfg["controls"]
        self.ptt_key = resolve_key(ctrl["push_to_talk_key"])
        self.pause_key = resolve_key(ctrl["pause_key"])
        self.reject_key = resolve_key(ctrl["reject_segment_key"])

    # ------------------------------------------------------------------
    # Audio-Callback + Worker
    # ------------------------------------------------------------------

    def audio_callback(self, indata, frames, time_info, status):
        """Sounddevice-Callback: schreibt immer in die Queue."""
        self.audio_queue.put(indata.copy())

    def recording_worker(self):
        """Hintergrund-Thread: sammelt Chunks + Live-Pegel."""
        while not self.finished and not self.aborted:
            try:
                chunk = self.audio_queue.get(timeout=0.1)
            except queue.Empty:
                continue
            with self.lock:
                if self.recording and not self.paused:
                    self.segment_chunks.append(chunk)
                    r = rms(chunk.flatten())
                    bar = level_bar(r)
                    print(f"\r  ● REC {bar}", end="", flush=True)

    def drain_queue(self):
        while not self.audio_queue.empty():
            try:
                self.audio_queue.get_nowait()
            except queue.Empty:
                break

    # ------------------------------------------------------------------
    # Segment-Verarbeitung
    # ------------------------------------------------------------------

    def process_segment(self) -> str:
        with self.lock:
            if not self.segment_chunks:
                return ""
            audio_data = np.concatenate(self.segment_chunks, axis=0).flatten()
            self.segment_chunks.clear()

        if audio_data.size == 0:
            return ""

        threshold = self.cfg["audio"]["silence_threshold_rms"]
        if rms(audio_data) < threshold:
            print("  (Stille – übersprungen)")
            return ""

        # float32 normalisieren
        audio_f32 = audio_data.astype(np.float32)
        if audio_f32.max() > 1.0:
            audio_f32 /= 32768.0

        duration = len(audio_f32) / self.sample_rate
        print(f"\n  Transkribiere ({duration:.1f}s) …", flush=True)
        text = self.transcriber.transcribe(audio_f32)

        if not text:
            print("  (kein Text erkannt)")
            return ""

        print(f"  -> {text}")
        return text

    def _process_and_store(self):
        if not self._processing_lock.acquire(blocking=False):
            print("\n  (Verarbeitung läuft noch – bitte warten)")
            return
        try:
            text = self.process_segment()
            if text:
                with self.lock:
                    self.text_blocks.append(text)
                if self.cfg["controls"]["segment_review"]:
                    print("  ✓ Übernommen (F9 zum Verwerfen)")
                else:
                    print("  ✓ Gespeichert")
                self._print_controls()
        finally:
            self._processing_lock.release()

    # ------------------------------------------------------------------
    # Anzeige
    # ------------------------------------------------------------------

    def _print_controls(self):
        ptt = self.cfg["controls"]["push_to_talk_key"].upper()
        model = self.cfg["correction"]["model"]
        n = len(self.text_blocks)
        print(f"  ┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄")
        print(f"  {n} Segment{'e' if n != 1 else ''} │ Ollama: {model} │ "
              f"[{ptt}]=Aufnahme  [F9]=Verwerfen  [F10]=Pause  [ESC]=Beenden")

    # ------------------------------------------------------------------
    # Tastatur-Handler
    # ------------------------------------------------------------------

    def on_key_press(self, key):
        if key == self.ptt_key:
            with self.lock:
                if self.paused:
                    return
                if not self.recording:
                    self.recording = True
                    self.drain_queue()
                    self.segment_chunks.clear()
                    print("\r  ● REC ", end="", flush=True)

    def on_key_release(self, key):
        if key == self.ptt_key:
            with self.lock:
                if not self.recording:
                    return
                self.recording = False
            print("\r  ■ STOPP" + " " * 50)
            threading.Thread(target=self._process_and_store, daemon=True).start()

        elif key == self.reject_key:
            with self.lock:
                if self.text_blocks:
                    removed = self.text_blocks.pop()
                    short = removed[:80] + "…" if len(removed) > 80 else removed
                    print(f"\n  ✗ Verworfen: {short}")
                else:
                    print("\n  (Keine Segmente zum Verwerfen)")
            self._print_controls()

        elif key == self.pause_key:
            with self.lock:
                self.paused = not self.paused
                status = "PAUSIERT" if self.paused else "FORTGESETZT"
            print(f"\n  [{status}]")
            self._print_controls()

        elif key == keyboard.Key.esc:
            print("\n\n  [ESC] Diktat beenden …")
            self.finished = True
            return False

    # ------------------------------------------------------------------
    # Hauptschleife
    # ------------------------------------------------------------------

    def run(self, device_idx: int):
        ptt_name = self.cfg["controls"]["push_to_talk_key"].upper()
        print(f"\n  Bereit – halte [{ptt_name}] zum Sprechen.\n")
        self._print_controls()

        stream = sd.InputStream(
            samplerate=self.sample_rate,
            channels=self.channels,
            blocksize=self.blocksize,
            device=device_idx,
            dtype="float32",
            callback=self.audio_callback,
        )

        worker = threading.Thread(target=self.recording_worker, daemon=True)
        worker.start()

        with stream:
            with keyboard.Listener(
                on_press=self.on_key_press,
                on_release=self.on_key_release,
            ) as listener:
                listener.join()

        worker.join(timeout=2.0)

        if self.aborted:
            print("\n  Diktat abgebrochen.")
            return

        with self.lock:
            blocks = list(self.text_blocks)

        if not blocks:
            print("\n  Kein Text aufgenommen.")
            return

        raw_text = "\n\n".join(blocks)

        print(f"\n{'=' * 60}")
        print(f"  ROHTEXT ({len(blocks)} Segment{'e' if len(blocks) != 1 else ''})")
        print(f"{'–' * 60}")
        print(raw_text[:800])
        if len(raw_text) > 800:
            print("  …")

        # Korrektur
        corrected = self._correct(raw_text)

        print(f"\n{'=' * 60}")
        print("  KORRIGIERTER TEXT")
        print(f"{'–' * 60}")
        print(corrected)
        print(f"{'=' * 60}")

        # Zwischenablage
        try:
            pyperclip.copy(corrected)
            print("  -> In Zwischenablage kopiert")
        except Exception:
            pass

        # Speichern
        self._save(corrected)

    # ------------------------------------------------------------------
    # Korrektur via Ollama
    # ------------------------------------------------------------------

    def _correct(self, text: str) -> str:
        cfg = self.cfg["correction"]
        base_url = cfg["ollama_base_url"]
        model = cfg["model"]
        temp = cfg.get("temperature", 0.1)
        max_tokens = cfg.get("max_tokens", 4096)
        paragraph_wise = cfg.get("paragraph_wise", True)

        print(f"\n  Korrektur mit Ollama ({model}):")

        if paragraph_wise:
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            corrected = []
            prev_context = ""
            for i, para in enumerate(paragraphs):
                print(f"    Absatz {i+1}/{len(paragraphs)} …", end=" ", flush=True)
                if prev_context:
                    prompt = CORRECTION_PARAGRAPH_PREFIX.format(
                        context=prev_context[-500:], text=para
                    )
                else:
                    prompt = CORRECTION_FULL_PREFIX.format(text=para)
                result = self._call_with_retry(base_url, model, prompt, temp, max_tokens)
                corrected.append(result)
                prev_context = result
                print("OK")
            return "\n\n".join(corrected)
        else:
            print("    …", end=" ", flush=True)
            result = self._call_with_retry(
                base_url, model,
                CORRECTION_FULL_PREFIX.format(text=text),
                temp, max_tokens,
            )
            print("OK")
            return result

    def _call_with_retry(
        self,
        base_url: str,
        model: str,
        user_prompt: str,
        temp: float,
        max_tokens: int,
        max_retries: int = 3,
    ) -> str:
        for attempt in range(1, max_retries + 1):
            try:
                return _call_ollama(
                    base_url, model,
                    CORRECTION_SYSTEM_PROMPT, user_prompt,
                    temperature=temp, max_tokens=max_tokens,
                )
            except Exception as exc:
                if attempt < max_retries:
                    wait = 2 ** attempt
                    print(f"\n  [Versuch {attempt}/{max_retries} fehlgeschlagen: {exc}]"
                          f"\n  Warte {wait}s …", file=sys.stderr)
                    time.sleep(wait)
                else:
                    print(f"\n  [Fehler nach {max_retries} Versuchen: {exc}]",
                          file=sys.stderr)
                    return user_prompt
        return user_prompt

    # ------------------------------------------------------------------
    # Speichern
    # ------------------------------------------------------------------

    def _save(self, text: str):
        out_cfg = self.cfg["output"]
        out_dir = Path(out_cfg["directory"])
        out_dir.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fmt = out_cfg["format"].lower()

        if fmt == "docx" and _HAS_DOCX:
            path = out_dir / f"diktat_{ts}.docx"
            self._save_docx(text, path, out_cfg)
        else:
            path = out_dir / f"diktat_{ts}.txt"
            path.write_text(text, encoding="utf-8")

        print(f"  Gespeichert: {path.resolve()}")

    def _save_docx(self, text: str, path: Path, cfg: dict):
        template = cfg.get("docx_template")
        if template and Path(template).is_file():
            doc = Document(template)
        else:
            doc = Document()
        font_name = cfg.get("docx_font", "Arial")
        font_size = cfg.get("docx_font_size_pt", 11)
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block == "[SEITENUMBRUCH]":
                doc.add_page_break()
                continue
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.font.name = font_name
            run.font.size = Pt(font_size)
        doc.save(path)


# =========================================================================
# Hauptprogramm
# =========================================================================

def main():
    if not _HAS_FASTER_WHISPER:
        sys.exit(
            "FEHLER: faster-whisper nicht installiert.\n"
            "  pip install faster-whisper"
        )

    cfg = load_config()
    base_url = cfg["correction"]["ollama_base_url"]

    print("=" * 60)
    print("  KI-Diktat-Ollama – Vollständig lokale Lösung")
    print("=" * 60)

    if not _check_ollama(base_url):
        sys.exit(
            f"FEHLER: Ollama nicht erreichbar ({base_url})\n"
            "  Starte Ollama: ollama serve"
        )

    available = _list_ollama_models(base_url)
    print(f"\n  Ollama erreichbar: {base_url}")
    if available:
        print(f"  Verfügbare Modelle: {', '.join(available)}")

    current_model = cfg["correction"]["model"]
    if available and current_model not in available:
        print(f"\n  [Warnung] Modell '{current_model}' nicht verfügbar.")
        print(f"  Tipp: ollama pull {current_model}")

    t_cfg = cfg["transcription"]
    print(f"\n  Transkription: FasterWhisper '{t_cfg['model']}' "
          f"({t_cfg['device']}/{t_cfg['compute_type']})")
    print(f"  Korrektur:     Ollama ({current_model})")

    try:
        device_idx = choose_input_device()
    except (RuntimeError, ValueError) as exc:
        sys.exit(f"FEHLER: {exc}")

    dev_info = sd.query_devices(device_idx)
    print(f"\n  Verwende: [{device_idx}] {dev_info['name']}")

    # Transcriber einmal erstellen – Modell wird beim ersten Diktat geladen
    transcriber = Transcriber(cfg)

    diktat_nr = 0
    while True:
        diktat_nr += 1
        model = cfg["correction"]["model"]

        if diktat_nr > 1:
            print(f"\n{'=' * 60}")
            print(f"  Neues Diktat #{diktat_nr}  (Ollama: {model})")
            print(f"{'=' * 60}")

        session = DictationSession(cfg, transcriber)
        session.run(device_idx)

        # Menü nach Diktat
        while True:
            model = cfg["correction"]["model"]
            print()
            print("  ┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄")
            print(f"  Korrektur-KI: Ollama ({model})")
            print("  ┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄┄")
            print("  [Enter]     Neues Diktat")
            print("  [B + Enter] Modell wechseln  "
                  "(mistral-small <-> gemma3)")
            print("  [Q + Enter] Tool beenden")
            try:
                choice = input("  > ").strip().lower()
            except (EOFError, KeyboardInterrupt):
                choice = "q"

            if choice in ("b", "modell", "wechseln"):
                order = AVAILABLE_MODELS
                current = cfg["correction"]["model"]
                next_idx = (order.index(current) + 1) % len(order) if current in order else 0
                next_model = order[next_idx]
                if available and next_model not in available:
                    print(f"\n  [!] '{next_model}' nicht in Ollama – "
                          f"Tipp: ollama pull {next_model}")
                else:
                    cfg["correction"]["model"] = next_model
                    print(f"\n  -> Gewechselt: Ollama ({next_model})")
                continue
            else:
                break

        if choice in ("q", "quit", "exit", "beenden"):
            break

    print("\n  Diktiertool beendet. Auf Wiedersehen!")


if __name__ == "__main__":
    main()
